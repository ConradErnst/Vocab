import requests
import json
import docx

#Function to retreive data
#Requests function to the Oxford API and returns json text
def retreive_text(n):
    app_id = 'APP_ID'
    app_key = 'APP_KEY'
    language = 'en-gb'
    for word in n:
        word_id = str(word)
        fields = 'definitions'
        url = 'https://od-api.oxforddictionaries.com/api/v2/entries/'  + language + '/' + word_id.lower() + '?fields=' + fields
        r = requests.get(url, headers = {'app_id' : app_id, 'app_key' : app_key})
        return(str(r.text))



#Function to clean text
#Takes in the json text and returns the word and the definition of the word
def clean_text(retreive_text):
    text = retreive_text
    dictionary_of_text = json.loads(text)
    definition = (dictionary_of_text['results'][0]['lexicalEntries'][0]['entries'][0]['senses'][0]['definitions'][0])
    part_of_speech = (dictionary_of_text['results'][0]['lexicalEntries'][0]['lexicalCategory']['id'])
    word = dictionary_of_text['word']
    return(definition, part_of_speech, word)



# Function to add text to document
# Adds the word and definition to word document
def add_text(clean_text):
    doc = docx.Document('Vocab.docx')
    doc.add_heading(clean_text[2], level=1)
    doc.add_paragraph(clean_text[1])
    doc.add_paragraph(clean_text[0])
    doc.save('Vocab.docx')


#Main call function
def main():
    words = str(input("Enter the word or words you would like to add to your dictionary separated by commas. \n "))
    list = words.split(",")

    try:
        for word in list:
            text = clean_text(retreive_text([word]))
            print("Word:", text[2])
            print("Part of speech:", text[1])
            print("Definition:", text[0])
            save = str(input("Do you want to save this word to your file? "))
            if save.lower() == "yes":
                add_text(text)
            else:
                break

    except:
        print("The word you are searching for is not correctly typed or the tense is unknown, try a different spelling or tense of the word ")
main()
